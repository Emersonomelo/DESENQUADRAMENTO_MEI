'''Codigo feito para fazer os formularios de desenquadramento do MEI de forma automatica'''

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.shared import Inches
import openpyxl
import pandas
import json
import urllib.request
import openpyxl

Document = Document()

def dados_da_receita(cnpj):
    
    url = 'http://receitaws.com.br/v1/cnpj/{0}'.format(cnpj)
    opener = urllib.request.build_opener()
    opener.addheaders = [
        ('User-agent',
        " Mozilla/5.0 (Windows NT 6.2; WOW64; rv:39.0) Gecko/20100101 Firefox/39.0")]

    with opener.open(url) as fd:
        content = fd.read().decode()

    dic = json.loads(content)
    nome = dic["nome"]
    planilha = openpyxl.Workbook()
    planilha.create_sheet('Planilha1teste', 0)
    planilha1 = planilha['Planilha1teste']

    planilha1['a1'].value ="abertura"
    planilha1['a2'].value ="situacao"
    planilha1['a3'].value ="tipo"
    planilha1['a4'].value ="nome"
    planilha1['a5'].value ="fantasia"
    planilha1['a6'].value ="porte"
    planilha1['a7'].value ="natureza_juridica"
    planilha1['a8'].value ="atividade_principal"
    planilha1['a9'].value ="atividades_secundarias"
    planilha1['a10'].value ="logradouro"
    planilha1['a11'].value ="numero"
    planilha1['a12'].value ="complemento"
    planilha1['a13'].value ="municipio"
    planilha1['a14'].value ="bairro"
    planilha1['a15'].value ="uf"
    planilha1['a16'].value ="cep"
    planilha1['a17'].value ="email"
    planilha1['a18'].value ="telefone"
    planilha1['a19'].value ="data_situacao"
    planilha1['a20'].value ="cnpj"
    planilha1['a21'].value ="ultima_atualizacao"
    planilha1['a22'].value ="status"
    planilha1['a23'].value ="efr"
    planilha1['a24'].value ="motivo_situacao"
    planilha1['a25'].value ="situacao_especial"
    planilha1['a26'].value ="data_situacao_especial"
    planilha1['a27'].value ="capital_social"
    planilha1['a28'].value ="qsa"
    planilha1['a29'].value ="extra"
    planilha1['a30'].value ="billing"
    planilha1['a31'].value ="cidade nascimento"
    planilha1['a32'].value ="data_de_nascimento" 
    planilha1['a33'].value ="estado_civil"
    planilha1['a34'].value ="profissão"
    planilha1['a35'].value ="rg"      
    planilha1['a36'].value ="expedido"
    planilha1['a37'].value ="CPF" 
    planilha1['a38'].value ="endereco"      
    planilha1['a39'].value ="mes" 
    planilha1['a40'].value ="ano" 
    planilha1['a41'].value ="data_desenquadramento"
    

    planilha1['b1'].value = str(dic["abertura"])
    planilha1['b2'].value = str(dic["situacao"])
    planilha1['b3'].value = str(dic["tipo"])
    planilha1['b4'].value = str(dic["nome"])
    planilha1['b5'].value = str(dic["fantasia"])
    planilha1['b6'].value = str(dic["porte"])
    planilha1['b7'].value = str(dic["natureza_juridica"])
    planilha1['b8'].value = str(dic["atividade_principal"][0]["text"])
    planilha1['b9'].value = str(dic["atividades_secundarias"])
    planilha1['b10'].value = str(dic["logradouro"])
    planilha1['b11'].value = str(dic["numero"])
    planilha1['b12'].value = str(dic["complemento"])
    planilha1['b13'].value = str(dic["municipio"])
    planilha1['b14'].value = str(dic["bairro"])
    planilha1['b15'].value = str(dic["uf"])
    planilha1['b16'].value = str(dic["cep"])
    planilha1['b17'].value = str(dic["email"])
    planilha1['b18'].value = str(dic["telefone"])
    planilha1['b19'].value = str(dic["data_situacao"])
    planilha1['b20'].value = str(dic["cnpj"])
    planilha1['b21'].value = str(dic["ultima_atualizacao"])
    planilha1['b22'].value = str(dic["status"])
    planilha1['b23'].value = str(dic["efr"])
    planilha1['b24'].value = str(dic["motivo_situacao"])
    planilha1['b25'].value = str(dic["situacao_especial"])
    planilha1['b26'].value = str(dic["data_situacao_especial"])
    planilha1['b27'].value = str(dic["capital_social"])
    planilha1['b28'].value = str(dic["qsa"])
    planilha1['b29'].value = str(dic["extra"])
    planilha1['b30'].value = str(dic["billing"])
    planilha1['b31'].value =input('Qual cidade de nascimento? ')
    planilha1['b32'].value =input('Qual a data de nascimento? ')
    planilha1['b33'].value =input('Qual estado civil? ')
    planilha1['b34'].value ='Empresário'
    planilha1['b35'].value =input('Qual RG? ')
    planilha1['b36'].value ='SSP/SP'
    cpf = str(dic["nome"])
    cpf = cpf[-11:]
    cpf = f'{cpf[0:3]}.{cpf[3:6]}.{cpf[6:9]}-{cpf[9:11]}'
        
    planilha1['b37'].value = cpf
    rua = planilha1['b10'].value
    numero = planilha1['b11'].value
    complemento = planilha1['b12'].value
    municipio = planilha1['b13'].value
    bairro = planilha1['b14'].value
    uf = planilha1['b15'].value
    
    cep = str(dic["cep"])
    cep = f'{cep[0:3]}{cep[3:]}'
    
    planilha1['b38'].value = f'{rua}, {numero}, {complemento}, Bairro: {bairro} - {municipio}, {uf} - CEP: {cep},'
    planilha1['b39'].value =input('Qual mes do desenquadramento? ') 
    planilha1['b40'].value =input('Qual ano do desenquadramento? ') 
    planilha1['b41'].value =str(input('Qual data de desenquadramento feito no site? '))
    
    planilha.save(f'{nome}.xlsx')
    return nome
        
    
def formulario(nome):
    
    #margens da pagina
    sections = Document.sections
    section = sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(0.5)

    #dicionario
    dados1 = f'C:\PYTHON_DADOS_RECEITA\{nome}.xlsx' 
    dados = openpyxl.load_workbook(dados1)

    nome_planilhas = dados.sheetnames
    planilha1 = dados['Planilha1teste']

    nome = planilha1['b4'].value
    cidade_nascimento = planilha1['b31'].value
    data_de_nascimento = pandas.to_datetime(planilha1['b32'].value).date()
    data_de_nascimento = data_de_nascimento.strftime('%d/%m/%Y')

    estado_civil = planilha1['b33'].value
    profissão = planilha1['b34'].value
    rg = planilha1['b35'].value
    expedido = planilha1['b36'].value
    CPF = planilha1['b37'].value
    endereco = planilha1['b38'].value
    mes = planilha1['b39'].value
    ano = planilha1['b40'].value
    cnpj = planilha1['b20'].value
    data_desenquadramento = pandas.to_datetime(planilha1['b41'].value).date()
    data_desenquadramento = data_desenquadramento.strftime('%d/%m/%Y')

    # Estilos:
    styles = Document.styles

    # Estilo do titulo
    titulo = styles.add_style("Titulo", WD_STYLE_TYPE.PARAGRAPH)
    titulo.font.name = "Arial"
    titulo.font.size = Pt(12)
    titulo.font.bold = True

    # Estilo do paragrafo
    paragrafo = styles.add_style("Paragraph", WD_STYLE_TYPE.PARAGRAPH)
    paragrafo.font.name = "Arial"
    paragrafo.font.size = Pt(12)
    paragrafo.font.bold = False

    # Estilo do paragrafo "Verdana"
    paragrafo = styles.add_style("Paragraph4", WD_STYLE_TYPE.PARAGRAPH)
    paragrafo.font.name = "Verdana"
    paragrafo.font.size = Pt(12)
    paragrafo.font.bold = False

    # Estilo do paragrafo com negrito
    paragrafo2 = styles.add_style("Paragraph2", WD_STYLE_TYPE.PARAGRAPH)
    paragrafo2.font.name = "Arial"
    paragrafo2.font.size = Pt(12)
    paragrafo2.font.bold = True

    T1 = Document.add_paragraph("DESENQUADRAMENTO", style="Titulo")
    T1.alignment = 1


    p1 = Document.add_paragraph(f"\n\tEu, {nome[:-12]}, nacionalidade Brasileira, natural de {cidade_nascimento}, nascido (a) em {data_de_nascimento}, estado civil: {estado_civil}, profissão {profissão}, portador (a) do documento de identidade RG nº {rg}, expedido (a) pelo(a) {expedido}, CPF nº {CPF}, residente à {endereco} na condição de titular, venho por meio deste requerer o DESENQUADRAMENTO do Microempreendedor Individual – MEI, efetuado no Portal do Empreendedor e não transmitido para o banco de dados da Jucesp, nos termos do Ofício Circular nº 43/2015/DREI/SRS/SMPE-PR do Departamento de Registro Empresarial e Integração/DREI, apresentando para tanto os documentos abaixo:", style="Paragraph")
    p1.alignment = 3

    p2 = Document.add_paragraph(f"\n•\tComprovante de desenquadramento emitido no Portal do Empreendedor.", style="Paragraph2")
    p2.alignment = 3

    p3 = Document.add_paragraph(f'\n\tFirmo a presente declaração sob as penas da lei (art. 1º. Da lei 7.115 de 29/08/1983), para que produza os efeitos legais, ciente de que, se comprovadamente falsa a declaração, sujeitar-me-ei, na qualidade de declarante às sanções civis, administrativas e criminais previstas na legislação aplicável. E será nulo de pleno direito perante o registro do comércio o ato a que se integre esta declaração, sem prejuízo das sanções a que estiver sujeito.', style="Paragraph")
    p3.alignment = 3

    p4 = Document.add_paragraph(f'\n\nSão Paulo, 01 de {mes} de {ano}',style="Paragraph")
    p4.alignment = 1

    p5 = Document.add_paragraph(f'\n\n_____________________________',style="Paragraph")
    p5.alignment = 1

    p6 = Document.add_paragraph(f'Nome e CPF',style="Paragraph")
    p6.alignment = 1

    p7 = Document.add_paragraph(f'\nObs.\n\tSe a empresa não estiver cadastrada na Junta Comercial deverá ser apresentado os seguintes documentos: (1) Certificado da Condição de Microempreendedor Individual – CCMEI; (2) Cadesp – extrato completo (se a atividade exigir) ou comprovante de residência do Microempreendedor Individual ou ainda, declaração escrita que conste o endereço residencial; (3) Cópia do RG ou outro documento de identificação pessoal que conste a data de nascimento; (4) Cópia do cartão do CNPJ.',style="Paragraph")
    p7.alignment = 3

    #quebra de pagina 
    Document.add_page_break()

    p8 = Document.add_paragraph(f'\nSão Paulo, 01 de {mes} de {ano}',style="Paragraph4")
    p8.alignment = 3

    p9 = Document.add_paragraph(f'\n\nAo Senhor',style="Paragraph4")
    p9.alignment = 3

    p10 = Document.add_paragraph(f'Presidente da Junta Comercial do Estado de São Paulo',style="Paragraph4")
    p10.alignment = 3

    p11 = Document.add_paragraph(f'\nAssunto: Desenquadramento MEI',style="Paragraph4")
    p11.alignment = 3

    p12 = Document.add_paragraph(f'\nSenhor Presidente,',style="Paragraph4")
    p12.alignment = 3

    p13 = Document.add_paragraph(f'\n\tNa qualidade de Titular da empresa {nome}, registrada sob CNPJ {cnpj}, solicito o desenquadramento da situação MEI perante a Junta Comercial do Estado de São Paulo.',style="Paragraph4")
    p13.alignment = 3

    p14 = Document.add_paragraph(f'\n\n\tInformo que referido desenquadramento foi realizado em {data_desenquadramento}, no portal do Empreendedor, conforme comprovante anexo.',style="Paragraph4")
    p14.alignment = 3

    p15 = Document.add_paragraph(f'\n\nAtenciosamente.',style="Paragraph4")
    p15.alignment = 3

    p16 = Document.add_paragraph(f'\n\n_____________________________________',style="Paragraph4")
    p16.alignment = 3


    Document.save(f'{nome}.docx')
    
    return nome
   

    
cnpj = input('Qual CNPJ: ')
x = dados_da_receita(cnpj)
y = formulario(x)
